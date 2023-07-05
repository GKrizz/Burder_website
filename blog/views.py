# from django.shortcuts import render

from django.shortcuts import render, redirect, get_object_or_404
from .models import User, Post
from django.shortcuts import render, redirect
from blog.forms import FileForm
from openpyxl import Workbook
from docx import Document
from django.http import HttpResponse
import openpyxl
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font
from django.http import HttpResponse
from django.utils import timezone

def signup(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        # Create a new user instance and save it to the database
        user = User(username=username, password=password)
        user.save()
        return redirect('login')  # Redirect to the login page

    return render(request, 'signup.html')

def login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        # Retrieve the user from the database
        user = User.objects.get(username=username, password=password)
        request.session['user_id'] = user.id  # Store the user's ID in the session
        return redirect('blog')  # Redirect to the home page

    return render(request, 'login.html')

def blog(request):
    # Retrieve all posts from the database
    posts = Post.objects.all()
    return render(request, 'blog.html', {'posts': posts})

def add_post(request):
    if request.method == 'POST':
        title = request.POST['title']
        content = request.POST['content']
        user_id = request.session.get('user_id')
        author = User.objects.get(id=user_id)
        # Create a new post instance and save it to the database
        post = Post(title=title, content=content, author=author)
        post.save()
        return redirect('blog')  # Redirect to the home page

    return render(request, 'add_post.html')

def edit_post(request, post_id):
    post = get_object_or_404(Post, id=post_id)
    if request.method == 'POST':
        post.title = request.POST['title']
        post.content = request.POST['content']
        post.save()
        return redirect('blog')  # Redirect to the home page

    return render(request, 'edit_post.html', {'post': post})

def delete_post(request, post_id):
    post = get_object_or_404(Post, id=post_id)
    if request.method == 'POST':
        post.delete()
        return redirect('blog')  # Redirect to the home page

    return render(request, 'delete_post.html', {'post': post})

def save_file(request):
    if request.method == 'POST':
        form = FileForm(request.POST, request.FILES)
        if form.is_valid():
            file = form.cleaned_data['file']
            file_extension = file.name.split('.')[-1].lower()
            
            if file_extension == 'xlsx':
                # Create an Excel workbook and save the file
                workbook = Workbook()
                sheet = workbook.active
                sheet.title = 'Sheet 1'
                # Write data to the sheet if required
                # file_path = 'C:\Users\GOBALA KRISHNAN\OneDrive\Desktop\burger-blog backend-2\excel\file.xlsx'
                file_path = r'media/my_file.xlsx '

                workbook.save(file_path)
            elif file_extension == 'docx':
                # Create a Word document and save the file
                document = Document()
                # Write content to the document if required
                file_path = r'media/my_file.docx'
                # file_path = r'C:\\Users\\GOBALA KRISHNAN\\OneDrive\\Desktop\\burger-blog backend-2\\word\\file.docx'

                document.save(file_path)
            else:
                return render(request, 'error.html', {'message': 'Invalid file format.'})
            
            return render(request, 'success.html', {'message': 'File saved successfully.'})
    else:
        form = FileForm()
    
    return render(request, 'file_form.html', {'form': form})


def index(request):
    return render(request, 'index.html')


def export_to_word(request):
    posts = Post.objects.all()

    document = Document()

    document.add_heading('Blog Posts', level=1)

    for post in posts:
        document.add_heading(post.title, level=2)
        document.add_paragraph(f"By: {post.author.username}")
        document.add_paragraph(post.content)
        document.add_paragraph(f"Created At: {post.created_at}")

    # Save the document
    document.save('blog_posts.docx')

    # Return the file as a response
    with open('blog_posts.docx', 'rb') as docx_file:
        response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=blog_posts.docx'

    return response



def export_to_excel(request):
    posts = Post.objects.all()

    # Create a new workbook and select the active sheet
    workbook = openpyxl.Workbook()
    sheet = workbook.active

    # Set the column headers
    headers = ['Title', 'Author', 'Content', 'Created At']
    for col_num, header in enumerate(headers, 1):
        col_letter = get_column_letter(col_num)
        sheet[f"{col_letter}1"] = header
        sheet[f"{col_letter}1"].font = Font(bold=True)

    # Write the post data to the sheet
    for row_num, post in enumerate(posts, 2):
        sheet[f"A{row_num}"] = post.title
        sheet[f"B{row_num}"] = post.author.username
        sheet[f"C{row_num}"] = post.content

        # Convert timezone-aware datetime to naive datetime
        created_at = post.created_at.astimezone(timezone.utc).replace(tzinfo=None)
        sheet[f"D{row_num}"] = created_at

    # Set the response headers for file download
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=blog_posts.xlsx'

    # Save the workbook to the response
    workbook.save(response)

    return response
